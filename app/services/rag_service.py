import json
import math
from dataclasses import dataclass
from pathlib import Path
from typing import Any, BinaryIO

import httpx
from sqlalchemy import delete, desc, select
from sqlalchemy.orm import Session

from app.core.config import Settings, get_settings
from app.models import KnowledgeChunk, KnowledgeDocument
from app.schemas.rag import RAGQueryResponse, RAGSource


SUPPORTED_EXTENSIONS = frozenset({".txt", ".md", ".pdf", ".docx"})


class RAGError(ValueError):
    pass


class EmbeddingConfigError(RuntimeError):
    pass


class RAGLLMConfigError(RuntimeError):
    pass


@dataclass(frozen=True)
class RankedChunk:
    chunk: KnowledgeChunk
    score: float


class EmbeddingService:
    def __init__(self, settings: Settings | None = None):
        self.settings = settings or get_settings()

    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        base_url = self.settings.embedding_base_url.strip()
        api_key = self.settings.embedding_api_key.strip()
        model = self.settings.embedding_model.strip()
        if not base_url or not api_key or not model:
            if self.settings.rag_local_demo_mode:
                return [_local_text_embedding(text) for text in texts]
            raise EmbeddingConfigError("请配置 EMBEDDING_BASE_URL、EMBEDDING_API_KEY、EMBEDDING_MODEL")

        response = httpx.post(
            f"{base_url.rstrip('/')}/embeddings",
            headers={"Authorization": f"Bearer {api_key}"},
            json={"model": model, "input": texts},
            timeout=30,
        )
        response.raise_for_status()
        payload = response.json()
        data = sorted(payload["data"], key=lambda item: item.get("index", 0))
        return [item["embedding"] for item in data]


class RAGService:
    def __init__(
        self,
        settings: Settings | None = None,
        embedding_service: EmbeddingService | None = None,
    ):
        self.settings = settings or get_settings()
        self.embedding_service = embedding_service or EmbeddingService(self.settings)

    def upload_document(
        self,
        db: Session,
        *,
        filename: str,
        content_type: str | None,
        document_type: str,
        file: BinaryIO,
    ) -> KnowledgeDocument:
        content = self._extract_text(filename, file)
        chunks = self._split_text(content)
        embeddings = self.embedding_service.embed_texts(chunks)
        document = KnowledgeDocument(
            filename=filename,
            document_type=document_type,
            content_type=content_type,
            status="ready",
            chunk_count=len(chunks),
        )
        db.add(document)
        db.flush()

        for index, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            db.add(
                KnowledgeChunk(
                    document_id=document.id,
                    chunk_index=index,
                    content=chunk,
                    embedding_json=json.dumps(embedding),
                )
            )
        db.commit()
        db.refresh(document)
        return document

    def list_documents(self, db: Session) -> list[KnowledgeDocument]:
        return list(
            db.scalars(
                select(KnowledgeDocument).order_by(desc(KnowledgeDocument.created_at))
            )
        )

    def delete_document(self, db: Session, document_id: int) -> bool:
        result = db.execute(
            delete(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
        )
        db.commit()
        return bool(result.rowcount)

    def query(self, db: Session, *, question: str, top_k: int | None = None) -> RAGQueryResponse:
        query_embedding = self.embedding_service.embed_texts([question])[0]
        ranked_chunks = self._rank_chunks(db, query_embedding, top_k or self.settings.rag_top_k)
        answer = self._generate_answer(question, ranked_chunks)
        return RAGQueryResponse(
            question=question,
            answer=answer,
            sources=[
                RAGSource(
                    document_id=ranked.chunk.document.id,
                    filename=ranked.chunk.document.filename,
                    document_type=ranked.chunk.document.document_type,
                    chunk_id=ranked.chunk.id,
                    chunk_index=ranked.chunk.chunk_index,
                    content=ranked.chunk.content,
                    score=ranked.score,
                )
                for ranked in ranked_chunks
            ],
        )

    def _extract_text(self, filename: str, file: BinaryIO) -> str:
        extension = Path(filename).suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise RAGError(f"不支持的文件类型: {extension}")

        raw = file.read()
        if isinstance(raw, str):
            raw_bytes = raw.encode("utf-8")
        else:
            raw_bytes = raw

        if extension in {".txt", ".md"}:
            return raw_bytes.decode("utf-8", errors="ignore").strip()
        if extension == ".pdf":
            return _extract_pdf_text(raw_bytes)
        if extension == ".docx":
            return _extract_docx_text(raw_bytes)
        raise RAGError(f"不支持的文件类型: {extension}")

    def _split_text(self, text: str) -> list[str]:
        normalized = " ".join(text.split())
        if not normalized:
            raise RAGError("文档内容为空")

        chunk_size = max(100, self.settings.rag_chunk_size)
        overlap = min(max(0, self.settings.rag_chunk_overlap), chunk_size - 1)
        chunks: list[str] = []
        start = 0
        while start < len(normalized):
            end = min(len(normalized), start + chunk_size)
            chunks.append(normalized[start:end])
            if end == len(normalized):
                break
            start = end - overlap
        return chunks

    def _rank_chunks(
        self,
        db: Session,
        query_embedding: list[float],
        top_k: int,
    ) -> list[RankedChunk]:
        chunks = db.scalars(select(KnowledgeChunk)).all()
        ranked = [
            RankedChunk(
                chunk=chunk,
                score=_cosine_similarity(query_embedding, json.loads(chunk.embedding_json)),
            )
            for chunk in chunks
        ]
        ranked.sort(key=lambda item: item.score, reverse=True)
        return ranked[:top_k]

    def _generate_answer(self, question: str, chunks: list[RankedChunk]) -> str:
        base_url = self.settings.llm_base_url.strip()
        api_key = self.settings.llm_api_key.strip()
        model = self.settings.llm_model.strip()
        if not base_url or not api_key or not model:
            if self.settings.rag_local_demo_mode:
                return _generate_local_demo_answer(question, chunks)
            raise RAGLLMConfigError("请配置 LLM_BASE_URL、LLM_API_KEY、LLM_MODEL")

        context = "\n\n".join(
            f"[来源 {index + 1}: {ranked.chunk.document.filename}]\n{ranked.chunk.content}"
            for index, ranked in enumerate(chunks)
        )
        response = httpx.post(
            f"{base_url.rstrip('/')}/chat/completions",
            headers={"Authorization": f"Bearer {api_key}"},
            json={
                "model": model,
                "messages": [
                    {
                        "role": "system",
                        "content": (
                            "你是安防知识库助手。只能依据给定知识库片段回答；"
                            "如果片段不足以回答，要明确说明缺少依据。回答使用中文。"
                            "请严格按以下 Markdown 结构输出："
                            "## 结论\n用 1-2 句话直接回答。"
                            "## 处理步骤\n用编号列表给出 3-5 条可执行步骤。"
                            "## 引用依据\n用短句说明依据来自哪些知识库片段。"
                        ),
                    },
                    {
                        "role": "user",
                        "content": f"问题：{question}\n\n知识库片段：\n{context}",
                    },
                ],
                "temperature": 0,
            },
            timeout=30,
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]


def _cosine_similarity(left: list[float], right: list[float]) -> float:
    if not left or not right or len(left) != len(right):
        return 0.0
    dot = sum(a * b for a, b in zip(left, right))
    left_norm = math.sqrt(sum(value * value for value in left))
    right_norm = math.sqrt(sum(value * value for value in right))
    if left_norm == 0 or right_norm == 0:
        return 0.0
    return dot / (left_norm * right_norm)


def _local_text_embedding(text: str, dimensions: int = 256) -> list[float]:
    vector = [0.0] * dimensions
    normalized = text.lower()
    for index, character in enumerate(normalized):
        bucket = (ord(character) + index * 31) % dimensions
        vector[bucket] += 1.0
    return vector


def _generate_local_demo_answer(question: str, chunks: list[RankedChunk]) -> str:
    if not chunks:
        return (
            "## 结论\n"
            "本地演示模式：当前知识库没有可引用的内容，无法基于文档回答。\n\n"
            "## 处理步骤\n"
            "1. 先上传制度、SOP、设备说明或历史事件材料。\n"
            "2. 上传完成后重新提问，让系统基于入库片段检索。\n"
            "3. 如果仍无结果，检查文档内容是否为空或问题是否过于宽泛。\n\n"
            "## 引用依据\n"
            "- 当前没有可用引用来源。"
        )

    leading = chunks[0].chunk.content
    excerpt = leading[:240]
    if len(leading) > 240:
        excerpt += "..."
    filename = chunks[0].chunk.document.filename
    return (
        "## 结论\n"
        f"本地演示模式：针对“{question}”，最相关的知识库依据来自《{filename}》。"
        "建议先按制度要求完成身份核验、现场确认和事件升级。\n\n"
        "## 处理步骤\n"
        "1. 查看驾驶舱中的事件时间、地点、摄像头编号和人员信息。\n"
        "2. 结合门禁记录和现场画面确认事件是否真实发生。\n"
        "3. 如果涉及限制区域、连续异常或无法确认身份，应升级给值班主管。\n"
        "4. 处置结束后记录结论，并标注意外情况或误报原因。\n\n"
        "## 引用依据\n"
        f"- 来源：《{filename}》\n"
        f"- 命中片段：{excerpt}"
    )


def _extract_pdf_text(raw_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RAGError("PDF 解析依赖未安装: pypdf") from exc

    import io

    reader = PdfReader(io.BytesIO(raw_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def _extract_docx_text(raw_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RAGError("DOCX 解析依赖未安装: python-docx") from exc

    import io

    document = Document(io.BytesIO(raw_bytes))
    return "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
